import os
import docx
from flask import Flask, render_template, request, send_file, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from docx.enum.style import WD_STYLE_TYPE

app = Flask(__name__, template_folder='templates')
app.config['UPLOAD_FOLDER'] = './uploads'
app.config['CONVERTED_FOLDER'] = './converted'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)

# Define consonants, vowels, and independents
consonants = "क ख ग घ ङ च छ ज झ ञ ट ठ ड ढ ण त थ द ध न प फ ब भ म य र ल व श ष स ह d [k x ?k ³ p N t > ¥ V B M < .k r Fk n /k u i Q c Hk e ; j y o 'k \"k l g"
vowels = "ँंःऺऻ़ ऽािीुूृॄॅॆेैोौः़ ऽ ¡ a %ऺऻ़ · f k h q w `ॄ Wॆ s S ks kS %़ ·"
independents = "अ आ इ ई उ ऊ ऋ ए ऐ ओ औ v v k b b Z m Å _ , ,s vks vkS"

def load_conversion_map(file_path: str) -> dict:
    """Load conversion map from a file."""
    conversion_map = {}
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            lines = [line.strip() for line in file]
            for line in lines:
                if ' ' in line:
                    original_chars, _, mapped_chars = line.partition(' ')
                    conversion_map[original_chars] = mapped_chars
                else:
                    original_char, mapped_char = line
                    conversion_map[original_char] = mapped_char
    return conversion_map

def convert_text(text: str, conversion_map: dict) -> str:
    """Convert text using the provided conversion map."""
    sorted_conversion_map = sorted(conversion_map.items(), key=lambda x: len(x[0]), reverse=True)
    output_text = ""
    i = 0
    while i < len(text):
        char = text[i]
        if char in consonants:
            # Handle consonant conversion
            if i + 1 < len(text) and text[i + 1] in vowels:
                combined_char = char + text[i + 1]
                if combined_char in conversion_map:
                    output_text += conversion_map[combined_char]
                    i += 2
                else:
                    output_text += conversion_map[char]
                    i += 1
            else:
                output_text += conversion_map[char]
                i += 1
        elif char in vowels:
            # Handle vowel conversion
            if i + 1 < len(text) and text[i + 1] in consonants:
                combined_char = char + text[i + 1]
                if combined_char in conversion_map:
                    output_text += conversion_map[combined_char]
                    i += 2
                else:
                    output_text += conversion_map[char]
                    i += 1
            else:
                output_text += conversion_map[char]
                i += 1
        elif char in independents:
            # Handle independent conversion
            output_text += conversion_map[char]
            i += 1
        else:
            output_text += char
            i += 1
    return output_text

def apply_conversion(element, conversion_map: dict, target_font: str) -> None:
    """Apply conversion to an element (paragraph, header, footer, table, caption, or heading)."""
    if hasattr(element, 'text'):
        element.text = convert_text(element.text, conversion_map)
        for run in element.runs:
            run.font.name = target_font

def process_docx(file_path: str, conversion_type: str) -> None:
    """Process the uploaded .docx file and save the converted file."""
    try:
        doc = docx.Document(file_path)
        conversion_map_file = f"{conversion_type}_map.txt"
        conversion_map = load_conversion_map(conversion_map_file)
        target_font = "Kruti Dev 010" if conversion_type == "unicode_to_krutidev" else "Arial"

        # Apply conversion to paragraphs
        for paragraph in doc.paragraphs:
            apply_conversion(paragraph, conversion_map, target_font)

        # Apply conversion to headers
        for header in doc.headers:
            apply_conversion(header, conversion_map, target_font)

        # Apply conversion to footers
        for footer in doc.footers:
            apply_conversion(footer, conversion_map, target_font)

        # Apply conversion to tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    apply_conversion(cell, conversion_map, target_font)

        # Apply conversion to captions
        for caption in doc.part.captions:
            apply_conversion(caption, conversion_map, target_font)

        # Apply conversion to heading styles
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.startswith('Heading'):
                for paragraph in style.paragraph_format.element.paragraphs:
                    apply_conversion(paragraph, conversion_map, target_font)

        # Save the converted document
        output_file = os.path.join(app.config['CONVERTED_FOLDER'], 'converted_' + os.path.basename(file_path))
        try:
            doc.save(output_file)
        except Exception as e:
            print(f"Error saving output file: {e}")

    except Exception as e:
        print(f"Error occurred during conversion: {e}")

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files['file']
        if file:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(file_path)
            conversion_type = request.form['conversion_type']
            process_docx(file_path, conversion_type)
            return redirect(url_for('download_file', filename='converted_' + file.filename))
    return render_template('index.html')

@app.errorhandler(405)
def method_not_allowed(e):
    return jsonify({'error': 'Method not allowed'}), 405

@app.route('/download/<filename>')
def download_file(filename):
    filename = secure_filename(filename)
    return send_file(os.path.join(app.config['CONVERTED_FOLDER'], filename), as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)